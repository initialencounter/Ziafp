from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
import os
import win32com.client
import pythoncom
from docx import Document

app = FastAPI()

# 定义请求体的数据模型
class EditDocRequest(BaseModel):
    source_path: str
    save_dir: str
    project_no: str
    project_name: str
    is_965: bool
    is_power_bank: bool
    en_name: str = ''

class EditDocxRequest(BaseModel):
    source_path: str
    project_no: str
    date: str
    signature_img_path: str

def edit_doc_file(source_path, save_dir, project_no, project_name, is_965, is_power_bank, en_name=''):
    try:
        # 创建 Word 应用程序对象
        # 如果是wps 则改成word.Application 如果是office 则改成word.Application.16
        word = win32com.client.Dispatch("Word.Application") 
        word.Visible = False

        # 打开文档
        doc = word.Documents.Open(source_path)
        print(source_path)
        range = doc.Content
        
        if is_965:
            if is_power_bank:
                range.Paragraphs(1).Range.Delete()
                range.InsertBefore(f"{project_name.split(' ')[0]}{en_name}\n")
            # 删除设备
            third_paragraph = range.Paragraphs(3)
            third_paragraph.Range.Delete()
            fourth_paragraph = range.Paragraphs(3)
            fourth_paragraph.Range.Delete()

        range.InsertBefore(f"物品名称：{project_name}\n")
        range.InsertBefore(f"项目编号：{project_no}\n")
        
        # 设置第一行的行距为1.0
        first_paragraph = range.Paragraphs(1)
        first_paragraph.Format.LineSpacingRule = 0
        # 如果是wps 则改成word.LinesToPoints(1) 如果是office 则改成12
        first_paragraph.Format.LineSpacing = word.LinesToPoints(1.0) 
        first_paragraph.Alignment = 0

        # 设置第二行的行距为1.5
        second_paragraph = range.Paragraphs(2)
        second_paragraph.Format.LineSpacingRule = 0
        second_paragraph.Format.LineSpacing = word.LinesToPoints(1.5)
        second_paragraph.Alignment = 0
        
        # 设置第三行的行距为1.5
        third_paragraph = range.Paragraphs(3)
        third_paragraph.Format.LineSpacingRule = 0
        third_paragraph.Format.LineSpacing = word.LinesToPoints(1.5)
        third_paragraph.Alignment = 0
        save_path = os.path.join(save_dir, f"{project_no}.doc")
        # 保存文档
        doc.SaveAs(save_path)
        doc.Close()
        doc = None
        return save_path
    except pythoncom.com_error as e:
        print("COM error occurred:", e)
    finally:
        # 仅在没有打开文档时退出 Word
        if 'word' in locals() and hasattr(word, 'Quit'):
            word.Quit()
            word = None


def edit_docx_file(source_path, signature):
    try:
        # 打开文档
        doc = Document(source_path)
        # 替换UN38.3.3(f)和UN38.3.3(g)
        for table in doc.tables:
            # 遍历每个行
            for row in table.rows:
                # 遍历每个单元格
                for cell in row.cells:
                    # 遍历单元格中的每个段落
                    for paragraph in cell.paragraphs:
                        if "UN38.3.3(f)" in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                "UN38.3.3(f)",
                                "UN38.3.3.1(f)或/or\nUN38.3.3.2(d)"
                            )
                        if "UN38.3.3(g)" in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                "UN38.3.3(g)",
                                "UN38.3.3.1(f)或/or\nUN38.3.3.2(d)"
                            )
        # 替换标题
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "锂电池UN38.3试验概要" in run.text:
                    run.text = run.text.replace(
                        "锂电池UN38.3试验概要",
                        "锂电池/钠离子电池UN38.3试验概要"
                    )
                if "Lithium Battery Test Summary" in run.text:
                    run.text = run.text.replace(
                        "Lithium Battery Test Summary",
                        "Test Summary"
                    )
        doc.save(source_path)
    except Exception as e:
        print("Error occurred:", e)

@app.post("/edit-doc")
async def edit_doc(request: EditDocRequest):
    print(request)
    try:
        save_path = edit_doc_file(
            source_path=request.source_path,
            save_dir=request.save_dir,
            project_no=request.project_no,
            project_name=request.project_name,
            is_965=request.is_965,
            is_power_bank=request.is_power_bank,
            en_name=request.en_name
        )
        return {"message": "Document edited successfully", "save_path": save_path}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/edit-docx")
async def edit_docx(request: EditDocxRequest):
    try:
        edit_docx_file(request.source_path, request.signature_img_path)
        return {"message": "Document edited successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    

# 启动应用程序
# uvicorn main:app --port 25457

