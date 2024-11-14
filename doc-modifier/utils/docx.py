from docx import Document

def edit_docx_file(source_path, signature_img_path=None):
    try:
        # 打开文档
        doc = Document(source_path)
        # 替换UN38.3.3(f)和UN38.3.3(g)
        for table in doc.tables:
            # 遍历每个行
            for row in table.rows:
                # 遍历每个单元格
                for cell in row.cells:
                    # 遍历单元格中的每个段落
                    for paragraph in cell.paragraphs:
                        if "UN38.3.3(f)" in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                "UN38.3.3(f)",
                                "UN38.3.3.1(f)或/or\nUN38.3.3.2(d)"
                            )
                        if "UN38.3.3(g)" in paragraph.text:
                            paragraph.text = paragraph.text.replace(
                                "UN38.3.3(g)",
                                "UN38.3.3.1(f)或/or\nUN38.3.3.2(d)"
                            )
        # 替换标题
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "锂电池UN38.3试验概要" in run.text:
                    run.text = run.text.replace(
                        "锂电池UN38.3试验概要",
                        "锂电池/钠离子电池UN38.3试验概要"
                    )
                if "Lithium Battery Test Summary" in run.text:
                    run.text = run.text.replace(
                        "Lithium Battery Test Summary",
                        "Test Summary"
                    )
        if signature_img_path:
            doc = replace_last_image_in_docx(doc, signature_img_path)
        doc.save(source_path)
    except Exception as e:
        print("Error occurred:", e)



def replace_last_image_in_docx(doc, new_image_path):
    # 获取所有图片的关系
    rels = doc.part.rels
    
    # 找到所有图片的关系
    image_rels = [rel for rel in rels.values() if "image" in rel.target_ref]
    
    # 检查是否有图片
    if image_rels:
        # 获取最后一张图片的关系
        last_image_rel = image_rels[-1]
        
        # 替换最后一张图片
        with open(new_image_path, 'rb') as new_image_file:
            last_image_rel.target_part._blob = new_image_file.read()
    
 
    return doc

